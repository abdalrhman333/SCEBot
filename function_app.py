"""
SCE HR Policy Chatbot — Azure Functions + Bot Framework + Teams
===============================================================
Serverless Azure Function that handles Bot Framework messages from Teams.
No App Service needed — runs on Consumption plan.
"""

import os
import json
import logging
import azure.functions as func
from botbuilder.core import (
    BotFrameworkAdapter,
    BotFrameworkAdapterSettings,
    TurnContext,
)
from botbuilder.schema import Activity, ActivityTypes

import pandas as pd
from docx import Document as DocxDocument
from openai import AzureOpenAI

# ─────────────────────────────────────────────────────────────────────────────
# CONFIGURATION — all from Azure Function App Settings
# ─────────────────────────────────────────────────────────────────────────────

AZURE_OPENAI_ENDPOINT    = os.getenv("AZURE_OPENAI_ENDPOINT", "https://pbi-agent-resource.cognitiveservices.azure.com/")
AZURE_OPENAI_API_KEY     = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION", "2025-04-01-preview")
DEPLOYMENT_NAME          = os.getenv("DEPLOYMENT_NAME", "gpt-5.4-mini")

BOT_APP_ID       = os.getenv("MicrosoftAppId", "")
BOT_APP_PASSWORD = os.getenv("MicrosoftAppPassword", "")

# Files stored in the Function App's file system or Azure Blob
# For Functions: place in a 'data' folder at the function app root
POLICIES_FOLDER     = os.getenv("POLICIES_FOLDER", os.path.join(os.path.dirname(__file__), "..", "data", "Policies"))
EMPLOYEE_TABLE_PATH = os.getenv("EMPLOYEE_TABLE_PATH", os.path.join(os.path.dirname(__file__), "..", "data", "employee_roles.xlsx"))

ALL_POLICY_FILES = [
    "_سياسات الموارد البشرية.docx",
    "_قواعد عمل صندوق الموظفين الاجتماعيInternal.docx",
    "_مصفوفة الصلاحياتInternal.docx",
    "لائحة تنظيم العمل الداخلي.docx",
]

ROLE_POLICY_FILES = {
    "HR": ALL_POLICY_FILES,
    "Employee": [
        "_سياسات الموارد البشرية.docx",
        "لائحة تنظيم العمل الداخلي.docx",
    ],
}

MAX_TOKENS        = 2500
TEMPERATURE       = 0.3
MAX_HISTORY_TURNS = 10

# ─────────────────────────────────────────────────────────────────────────────
# CLIENTS (initialized once per cold start, reused across invocations)
# ─────────────────────────────────────────────────────────────────────────────

openai_client = AzureOpenAI(
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
    api_key=AZURE_OPENAI_API_KEY,
    api_version=AZURE_OPENAI_API_VERSION,
)

adapter_settings = BotFrameworkAdapterSettings(
    app_id=BOT_APP_ID,
    app_password=BOT_APP_PASSWORD,
)
adapter = BotFrameworkAdapter(adapter_settings)


async def on_error(context: TurnContext, error: Exception):
    logging.error(f"Bot error: {error}")
    await context.send_activity("عذراً، حدث خطأ في المعالجة. يرجى المحاولة مرة أخرى.")

adapter.on_turn_error = on_error

# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT & ROLE LOADING (same as your Fabric notebook)
# ─────────────────────────────────────────────────────────────────────────────

def extract_text_from_docx(filepath: str) -> str:
    doc = DocxDocument(filepath)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                paragraphs.append(" | ".join(row_texts))
    return "\n\n".join(paragraphs)


def load_all_policies(base_path: str, filenames: list) -> dict:
    loaded = {}
    for name in filenames:
        loaded[name] = extract_text_from_docx(os.path.join(base_path, name))
    return loaded


def load_employee_roles(filepath: str) -> pd.DataFrame:
    df = pd.read_excel(filepath)
    df["email"] = df["email"].str.strip().str.lower()
    df["position_type"] = df["position_type"].str.strip()
    return df


# Pre-load at cold start (stays warm across invocations on same instance)
roles_df             = load_employee_roles(EMPLOYEE_TABLE_PATH)
policy_docs_hr       = load_all_policies(POLICIES_FOLDER, ROLE_POLICY_FILES["HR"])
policy_docs_employee = load_all_policies(POLICIES_FOLDER, ROLE_POLICY_FILES["Employee"])
logging.info(f"Loaded {len(roles_df)} employees, {len(policy_docs_hr)} HR docs, {len(policy_docs_employee)} Employee docs")

# ─────────────────────────────────────────────────────────────────────────────
# ROLE DETECTION
# ─────────────────────────────────────────────────────────────────────────────

def detect_user_role(user_email: str) -> dict:
    match = roles_df[roles_df["email"] == user_email.strip().lower()]
    if match.empty:
        return {
            "full_name": "موظف",
            "email": user_email,
            "department": "غير محدد",
            "position_type": "Employee",
            "allowed_files": ROLE_POLICY_FILES["Employee"],
        }
    row = match.iloc[0]
    role = row["position_type"]
    return {
        "full_name": row["full_name"],
        "email": row["email"],
        "department": row["department"],
        "position_type": role,
        "allowed_files": ROLE_POLICY_FILES.get(role, ROLE_POLICY_FILES["Employee"]),
    }

# ─────────────────────────────────────────────────────────────────────────────
# SYSTEM PROMPT BUILDER (identical to Fabric version)
# ─────────────────────────────────────────────────────────────────────────────

def build_system_prompt(policy_docs: dict, user_role_info: dict) -> str:
    docs_block_parts = []
    for idx, (filename, content) in enumerate(policy_docs.items(), start=1):
        doc_header = f"{'═'*70}\n  المستند {idx}: {filename}\n{'═'*70}"
        docs_block_parts.append(f"{doc_header}\n\n{content}")
    docs_block = "\n\n".join(docs_block_parts)
    role_label = "موظف موارد بشرية" if user_role_info["position_type"] == "HR" else "موظف عام"

    return f"""أنت مساعد ذكاء اصطناعي متخصص في سياسات الموارد البشرية للشركة.
المستخدم الحالي: {user_role_info['full_name']} — {role_label} — قسم: {user_role_info['department']}

مهمتك هي مساعدة الموظفين بالإجابة على أسئلتهم المتعلقة بالسياسات والإجراءات الداخلية للشركة.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
 قواعد السلوك والرد (يجب الالتزام بها في جميع الأحوال)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

1. **اللغة والأسلوب:**
   - استخدم اللغة العربية الفصحى الرسمية دائمًا في ردودك ما لم يطلب منك غير ذلك.
   - يمكنك فهم الأسئلة المكتوبة بالعامية أو الإنجليزية والرد عليها بالعربية الرسمية.
   - كن مهذبًا ومحترمًا وودودًا في جميع ردودك.

2. **الدقة والمصدر:**
   - أجب فقط استنادًا إلى محتوى وثائق السياسات المُدرجة أدناه.
   - إذا لم تجد إجابة واضحة في الوثائق، أخبر المستخدم بأدب واقترح عليه التواصل مع قسم الموارد البشرية مباشرةً.
   - لا تخترع أو تفترض معلومات غير موجودة في الوثائق.
   - في نهاية كل رد، أضف سطراً يبدأ بـ [المصدر:] يذكر فيه اسم الملف أو الملفات التي استندت إليها.

3. **هيكل الإجابة:**
   - اجعل إجاباتك واضحة ومنظمة ومختصرة.
   - استخدم القوائم والنقاط عند الإجابة على أسئلة متعددة الجوانب.
   - إذا كانت المعلومة تتطلب خطوات، رتّبها بشكل منطقي ومتسلسل.
   - لا تُضف أي اقتراحات أو عروض مساعدة إضافية في نهاية ردك. أجب على السؤال فقط وأنهِ ردك بالمصدر مباشرةً.

4. **الحياد والموضوعية:**
   - لا تُبدِ آراء شخصية حول السياسات.
   - إذا كانت السياسة تحتمل تفسيرين، اعرض الاثنين وأوضح أن التأكيد يعود لقسم الموارد البشرية.

5. **الخصوصية والسرية:**
   - لا تطلب من المستخدم بيانات شخصية حساسة.
   - تذكّر دائمًا أن المعلومات التي تُقدمها ذات طابع توجيهي عام، وأن القرار النهائي يعود لقسم الموارد البشرية.

6. **التحكم بالوصول:**
   - أجب فقط بناءً على الوثائق المُتاحة أدناه. لا تشير إلى وثائق أخرى قد تكون موجودة ولكن غير مُدرجة هنا.
   - إذا سأل المستخدم عن موضوع غير مشمول في الوثائق المتاحة له، أخبره بأدب أن هذه المعلومة غير متوفرة حالياً واقترح التواصل مع قسم الموارد البشرية.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
 وثائق سياسات الشركة (قاعدة المعرفة — مرجعك الوحيد)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

{docs_block}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
 ملاحظة ختامية
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
أنت تتحدث مع {role_label}. كُن دليله الموثوق ضمن نطاق الوثائق المتاحة فقط.
"""

# ─────────────────────────────────────────────────────────────────────────────
# CHATBOT CLASS (unchanged from Fabric version)
# ─────────────────────────────────────────────────────────────────────────────

class HRChatbot:
    def __init__(self, system_prompt: str):
        self.system_prompt = system_prompt
        self.history = []

    def ask(self, user_input: str) -> str:
        input_messages = [{"role": "system", "content": self.system_prompt}]
        input_messages.extend(self.history[-(MAX_HISTORY_TURNS * 2):])
        input_messages.append({"role": "user", "content": user_input})

        response = openai_client.responses.create(
            model=DEPLOYMENT_NAME,
            input=input_messages,
            temperature=TEMPERATURE,
            max_output_tokens=MAX_TOKENS,
        )
        reply = response.output_text.strip()
        self.history.append({"role": "user", "content": user_input})
        self.history.append({"role": "assistant", "content": reply})
        return reply

    def clear(self):
        self.history = []


# ─────────────────────────────────────────────────────────────────────────────
# IN-MEMORY SESSION STORE
# ─────────────────────────────────────────────────────────────────────────────
# Survives across invocations on same warm instance.
# For persistence across cold starts, swap for Cosmos DB (you already have it).

user_sessions: dict[str, HRChatbot] = {}


def get_or_create_session(user_id: str, user_email: str) -> HRChatbot:
    if user_id not in user_sessions:
        user_info = detect_user_role(user_email)
        docs = policy_docs_hr if user_info["position_type"] == "HR" else policy_docs_employee
        system_prompt = build_system_prompt(docs, user_info)
        user_sessions[user_id] = HRChatbot(system_prompt)
        logging.info(f"New session: {user_email} ({user_info['position_type']})")
    return user_sessions[user_id]


# ─────────────────────────────────────────────────────────────────────────────
# BOT MESSAGE HANDLER
# ─────────────────────────────────────────────────────────────────────────────

async def on_message(turn_context: TurnContext):
    if turn_context.activity.type != ActivityTypes.message:
        return

    user_text = turn_context.activity.text or ""
    user_id   = turn_context.activity.from_property.id

    # Teams provides the user's name; for email, see guide on Graph API lookup
    user_email = getattr(turn_context.activity.from_property, "name", "") or ""
    if not user_email or "@" not in user_email:
        user_email = f"{user_id}@unknown.com"

    # Commands
    if user_text.strip() in ["مسح", "مسح المحادثة", "/clear", "/reset"]:
        if user_id in user_sessions:
            user_sessions[user_id].clear()
            del user_sessions[user_id]
        await turn_context.send_activity("✅ تم مسح المحادثة. يمكنك البدء من جديد.")
        return

    if user_text.strip() in ["/help", "مساعدة"]:
        await turn_context.send_activity(
            "🤖 **مساعد سياسات الموارد البشرية**\n\n"
            "اكتب سؤالك حول سياسات الشركة وسأجيبك.\n\n"
            "**أمثلة:**\n"
            "- ما هي أنواع الإجازات المتاحة؟\n"
            "- ما هي ساعات الدوام الرسمي؟\n\n"
            "**أوامر:** `/clear` — مسح | `/help` — مساعدة"
        )
        return

    # Typing indicator
    await turn_context.send_activity(Activity(type=ActivityTypes.typing))

    # Get session and respond
    bot = get_or_create_session(user_id, user_email)
    try:
        reply = bot.ask(user_text)
        await turn_context.send_activity(reply)
    except Exception as e:
        logging.error(f"OpenAI error for {user_id}: {e}")
        await turn_context.send_activity("عذراً، حدث خطأ أثناء معالجة سؤالك. يرجى المحاولة مرة أخرى.")


# ─────────────────────────────────────────────────────────────────────────────
# AZURE FUNCTION ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────

app = func.FunctionApp(http_auth_level=func.AuthLevel.ANONYMOUS)


@app.route(route="api/messages", methods=["POST"])
async def messages(req: func.HttpRequest) -> func.HttpResponse:
    """Bot Framework messaging endpoint — called by Azure Bot Service."""

    if "application/json" not in req.headers.get("Content-Type", ""):
        return func.HttpResponse(status_code=415)

    body = req.get_json()
    activity = Activity().deserialize(body)
    auth_header = req.headers.get("Authorization", "")

    try:
        await adapter.process_activity(activity, auth_header, on_message)
        return func.HttpResponse(status_code=200)
    except Exception as e:
        logging.error(f"process_activity error: {e}")
        return func.HttpResponse(status_code=500)


@app.route(route="api/health", methods=["GET"])
async def health(req: func.HttpRequest) -> func.HttpResponse:
    return func.HttpResponse(
        json.dumps({"status": "healthy", "active_sessions": len(user_sessions)}),
        mimetype="application/json",
    )
